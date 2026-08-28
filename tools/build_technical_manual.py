"""Build the maintained technical manual Markdown source as a polished DOCX.

The Markdown file is the single content source. This converter intentionally
supports the subset used by the manual: headings, paragraphs, block quotes,
real bulleted/numbered lists, fenced code blocks, and pipe tables.
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "docs" / "TECHNICAL_MANUAL.md"
DEFAULT_OUTPUT = ROOT / "docs" / "创作者内容采集台技术手册.docx"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "606B78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "BCC7D3"
WHITE = "FFFFFF"
BLACK = "202124"


def set_run_font(run, *, latin="Calibri", east_asia="Microsoft YaHei", size=None,
                 color=None, bold=None, italic=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, *, before=0, after=0, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    fmt.widow_control = True


def shade(element, fill: str):
    shd = element.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        element.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_border(paragraph, *, color=BORDER, size="6", space="4"):
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM),
        ("bottom", CELL_MARGIN_TOP_BOTTOM),
        ("start", CELL_MARGIN_START_END),
        ("end", CELL_MARGIN_START_END),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths: list[int]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = table._tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        old_grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)
        borders.append(node)


def make_abstract_num(numbering, *, abstract_id: int, ordered: bool):
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        if ordered:
            lvl_text.set(qn("w:val"), f"%{level + 1}.")
        else:
            lvl_text.set(qn("w:val"), ("•", "◦", "▪")[level])
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(540 + level * 360))
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(540 + level * 360))
        ind.set(qn("w:hanging"), "270")
        ppr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        lvl.append(ppr)
        if not ordered:
            rpr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Arial")
            fonts.set(qn("w:hAnsi"), "Arial")
            rpr.append(fonts)
            lvl.append(rpr)
        abstract.append(lvl)
    numbering.append(abstract)


def make_num(numbering, *, num_id: int, abstract_id: int):
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)


def install_numbering(document: Document) -> tuple[int, int]:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(value) for value in (
            node.get(qn("w:abstractNumId")) for node in numbering.findall(qn("w:abstractNum"))
        ) if value is not None
    ]
    num_ids = [
        int(value) for value in (
            node.get(qn("w:numId")) for node in numbering.findall(qn("w:num"))
        ) if value is not None
    ]
    abstract_start = max(abstract_ids, default=0) + 1
    num_start = max(num_ids, default=0) + 1
    make_abstract_num(numbering, abstract_id=abstract_start, ordered=False)
    make_abstract_num(numbering, abstract_id=abstract_start + 1, ordered=True)
    make_num(numbering, num_id=num_start, abstract_id=abstract_start)
    make_num(numbering, num_id=num_start + 1, abstract_id=abstract_start + 1)
    return num_start, num_start + 1


def apply_numbering(paragraph, num_id: int, level: int = 0):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(max(0, min(level, 2))))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def configure_styles(document: Document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for name in ("List Paragraph",):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.25

    code = styles.add_style("Manual Code", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string("263238")
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.05

    quote = styles.add_style("Manual Note", 1)
    quote.font.name = "Calibri"
    quote.font.size = Pt(10.5)
    quote.font.color.rgb = RGBColor.from_string("334155")
    quote._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    quote.paragraph_format.left_indent = Inches(0.18)
    quote.paragraph_format.right_indent = Inches(0.12)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(8)
    quote.paragraph_format.line_spacing = 1.18


def configure_page(document: Document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_page_field(paragraph):
    paragraph.add_run("第 ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    for node in (begin, instr, separate, display, end):
        run.append(node)
    paragraph._p.append(run)
    paragraph.add_run(" 页")
    for item in paragraph.runs:
        set_run_font(item, size=8.5, color=MUTED)


def configure_header_footer(document: Document):
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "创作者内容采集台技术手册  |  运行、维护与持续更新"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(paragraph, after=0, line=1.0)
    for run in paragraph.runs:
        set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=0, line=1.0)
    add_page_field(paragraph)


def add_inline(paragraph, text: str, *, default_size=None):
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=default_size)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, latin="Consolas", size=(default_size or 10.5), color=DARK_BLUE)
            shade(run._element.get_or_add_rPr(), LIGHT_GRAY)
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, bold=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=default_size)


def add_cover(document: Document, title: str, metadata: list[str]):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(94)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(kicker, after=18, line=1.0)
    run = kicker.add_run("TECHNICAL MANUAL")
    set_run_font(run, size=10, color=BLUE, bold=True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=10, line=1.05)
    run = paragraph.add_run(title)
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=50, line=1.15)
    run = subtitle.add_run("运行、维护、排障与持续更新指南")
    set_run_font(run, size=14, color=DARK_BLUE)

    for item in metadata:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=4, line=1.0)
        clean = re.sub(r"\s{2,}", " ", item).strip()
        add_inline(paragraph, clean, default_size=10)
        for run in paragraph.runs:
            if run.font.color.rgb is None:
                run.font.color.rgb = RGBColor.from_string(MUTED)

    callout = document.add_paragraph()
    callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(callout, before=36, after=0, line=1.15)
    run = callout.add_run("当前正式入口：app.py  |  内容结构版本：16  |  最大并发：3")
    set_run_font(run, size=9.5, color=WHITE, bold=True)
    shade(callout._p.get_or_add_pPr(), BLUE)
    callout.paragraph_format.left_indent = Inches(0.55)
    callout.paragraph_format.right_indent = Inches(0.55)

    document.add_page_break()


def split_table_row(line: str) -> list[str]:
    value = line.strip().strip("|")
    return [cell.strip() for cell in value.split("|")]


def is_separator_row(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def allocate_widths(rows: list[list[str]]) -> list[int]:
    columns = len(rows[0])
    scores = []
    for index in range(columns):
        lengths = [len(row[index]) if index < len(row) else 0 for row in rows]
        score = min(max(max(lengths, default=1), 6), 50)
        scores.append(score)
    minimum = 900 if columns >= 4 else 1100
    widths = [minimum] * columns
    remaining = CONTENT_WIDTH_DXA - sum(widths)
    score_sum = sum(scores) or columns
    for index, score in enumerate(scores):
        widths[index] += int(remaining * score / score_sum)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(document: Document, rows: list[list[str]]):
    if not rows:
        return
    columns = max(len(row) for row in rows)
    normalized = [row + [""] * (columns - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=columns)
    set_table_geometry(table, allocate_widths(normalized))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for r_index, row in enumerate(normalized):
        for c_index, value in enumerate(row):
            cell = table.cell(r_index, c_index)
            if r_index == 0:
                shade(cell._tc.get_or_add_tcPr(), LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line=1.12)
            if r_index == 0 or (columns >= 3 and c_index == 0):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(paragraph, value, default_size=9.2)
            if r_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_code_block(document: Document, lines: Iterable[str]):
    values = list(lines) or [""]
    for index, value in enumerate(values):
        paragraph = document.add_paragraph(style="Manual Code")
        add_inline(paragraph, value.replace("\t", "    "), default_size=8.5)
        shade(paragraph._p.get_or_add_pPr(), LIGHT_GRAY)
        paragraph.paragraph_format.left_indent = Inches(0.14)
        paragraph.paragraph_format.right_indent = Inches(0.10)
        if index == 0:
            paragraph.paragraph_format.space_before = Pt(5)
        if index == len(values) - 1:
            paragraph.paragraph_format.space_after = Pt(7)


def add_note(document: Document, text: str):
    paragraph = document.add_paragraph(style="Manual Note")
    add_inline(paragraph, text, default_size=10.5)
    shade(paragraph._p.get_or_add_pPr(), LIGHT_GRAY)
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    borders.append(left)


def build(source: Path, output: Path):
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_page(document)
    configure_styles(document)
    configure_header_footer(document)
    bullet_num, decimal_num = install_numbering(document)

    title = lines[0][2:].strip() if lines and lines[0].startswith("# ") else source.stem
    metadata = []
    index = 1
    while index < len(lines):
        line = lines[index]
        if line.startswith(">"):
            metadata.append(line.lstrip("> ").rstrip())
        elif metadata and not line.strip():
            index += 1
            break
        index += 1
    add_cover(document, title, metadata)

    paragraph_buffer: list[str] = []

    def flush_paragraph():
        if not paragraph_buffer:
            return
        text = " ".join(item.strip() for item in paragraph_buffer).strip()
        paragraph_buffer.clear()
        if text:
            paragraph = document.add_paragraph()
            set_paragraph_spacing(paragraph, after=6, line=1.25)
            add_inline(paragraph, text)

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(document, code_lines)
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and is_separator_row(lines[index + 1]):
            flush_paragraph()
            rows = [split_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_table_row(lines[index]))
                index += 1
            add_table(document, rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1)) - 1
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_inline(paragraph, heading.group(2))
            index += 1
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            add_note(document, stripped.lstrip("> "))
            index += 1
            continue
        bullet = re.match(r"^(\s*)-\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            level = min(len(bullet.group(1).replace("\t", "    ")) // 2, 2)
            text = bullet.group(2)
            if text.startswith("[ ] "):
                text = "☐ " + text[4:]
            elif text.startswith("[x] ") or text.startswith("[X] "):
                text = "☒ " + text[4:]
            paragraph = document.add_paragraph(style="List Paragraph")
            apply_numbering(paragraph, bullet_num, level)
            add_inline(paragraph, text)
            index += 1
            continue
        ordered = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if ordered:
            flush_paragraph()
            level = min(len(ordered.group(1).replace("\t", "    ")) // 2, 2)
            paragraph = document.add_paragraph(style="List Paragraph")
            apply_numbering(paragraph, decimal_num, level)
            add_inline(paragraph, ordered.group(2))
            index += 1
            continue
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        paragraph_buffer.append(line)
        index += 1
    flush_paragraph()

    # Apply quiet rules to Heading 1 paragraphs after all content exists.
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Heading 1":
            paragraph_border(paragraph, color="D8E0E8", size="4", space="5")

    core = document.core_properties
    core.title = title
    core.subject = "技术人员通用的运行、维护、排障与持续更新手册"
    core.author = "技术团队"
    core.keywords = "技术手册, Flask, 采集, OCR, AI, 运维"
    core.comments = "由 docs/TECHNICAL_MANUAL.md 自动生成；禁止只编辑 DOCX。"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main():
    parser = argparse.ArgumentParser(description="从 Markdown 构建技术手册 DOCX")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
